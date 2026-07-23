"""Gera o documento v2 (um por utilizador) de recolha de vocabulario para a IA."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

CASTANHO = RGBColor(0x5A, 0x3E, 0x2B)
CASTANHO_MEDIO = RGBColor(0x8B, 0x6F, 0x4E)
CINZENTO = RGBColor(0x80, 0x80, 0x80)
AVISO = RGBColor(0x8A, 0x5A, 0x00)
BEGE = "EFE7DA"
CREME = "FBF8F3"
CREME_AVISO = "FDF6E7"

LARGURA_UTIL = Cm(17)


def sombrear(celula, cor):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), cor)
    celula._tc.get_or_add_tcPr().append(shd)


def borda_esquerda(celula, cor_hex):
    tcPr = celula._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for lado, tamanho in (("left", "24"), ("top", "4"), ("bottom", "4"), ("right", "4")):
        el = OxmlElement(f"w:{lado}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), tamanho)
        el.set(qn("w:color"), cor_hex)
        borders.append(el)
    tcPr.append(borders)


def escrever(p, texto, *, bold=False, italic=False, size=10.5, cor=None):
    run = p.add_run(texto)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if cor is not None:
        run.font.color.rgb = cor
    return run


def paragrafo(doc, texto="", *, bold=False, italic=False, size=10.5, cor=None, depois=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(depois)
    p.paragraph_format.space_before = Pt(0)
    if texto:
        escrever(p, texto, bold=bold, italic=italic, size=size, cor=cor)
    return p


def titulo(doc, texto, nivel=1):
    p = doc.add_paragraph(style="Heading 1" if nivel == 1 else "Heading 2")
    p.paragraph_format.space_before = Pt(16 if nivel == 1 else 12)
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(15 if nivel == 1 else 12)
    run.font.color.rgb = CASTANHO
    run.font.name = "Calibri"
    return p


def celula_texto(celula, texto, *, bold=False, italic=False, cor=None, size=10):
    celula.text = ""
    p = celula.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if texto:
        escrever(p, texto, bold=bold, italic=italic, size=size, cor=cor)


def tabela(doc, cabecalhos, larguras, exemplos=(), vazias=8, fixos=()):
    t = doc.add_table(rows=1, cols=len(cabecalhos))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False

    for i, texto in enumerate(cabecalhos):
        c = t.rows[0].cells[i]
        celula_texto(c, texto, bold=True, cor=CASTANHO)
        sombrear(c, BEGE)

    for linha in fixos:  # linhas de dados reais, a preto
        celulas = t.add_row().cells
        for i, texto in enumerate(linha):
            celula_texto(celulas[i], texto)

    for linha in exemplos:  # linhas de exemplo, a cinzento e italico
        celulas = t.add_row().cells
        for i, texto in enumerate(linha):
            celula_texto(celulas[i], texto, italic=True, cor=CINZENTO)

    for _ in range(vazias):
        for c in t.add_row().cells:
            celula_texto(c, "")

    for linha in t.rows:
        for i, c in enumerate(linha.cells):
            c.width = Cm(larguras[i])

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def caixa(doc, titulo_caixa, linhas, *, cor_borda="5A3E2B", fundo=CREME, cor_titulo=CASTANHO):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    c = t.rows[0].cells[0]
    c.width = LARGURA_UTIL
    sombrear(c, fundo)
    borda_esquerda(c, cor_borda)

    c.text = ""
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    escrever(p, titulo_caixa, bold=True, size=11, cor=cor_titulo)
    for texto in linhas:
        pp = c.add_paragraph()
        pp.paragraph_format.space_after = Pt(3)
        escrever(pp, texto)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# Nomes exatamente como estão gravados no Martelo (tabela producao/users).
RESPONSAVEIS_REAIS = [
    "Paulo", "Pedro", "Angela", "Bruno", "Dario", "Ana", "Marcia", "Andreia",
    "Catia", "Elisabete",
]


def gerar(destino: Path, nome_pessoa: str) -> None:
    doc = Document()

    s = doc.sections[0]
    s.page_width, s.page_height = Cm(21), Cm(29.7)
    for lado in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(s, lado, Cm(2))

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # ---------------- capa ----------------
    paragrafo(doc, "MARTELO ORÇAMENTOS V3", bold=True, size=10, cor=CASTANHO_MEDIO, depois=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    escrever(p, "Ensinar o assistente a responder", bold=True, size=20, cor=CASTANHO)
    paragrafo(
        doc,
        f"Perfil de linguagem e preferências — preenchido por {nome_pessoa}",
        size=11,
        cor=RGBColor(0x5A, 0x5A, 0x5A),
        depois=12,
    )

    caixa(
        doc,
        "Porque é que este documento é importante",
        [
            "O Martelo V3 vai passar a ter um assistente que responde a perguntas escritas "
            "como as dizemos no dia a dia — por exemplo «que roupeiros do Paulo estão "
            "atrasados?» — sem obrigar a mexer nos filtros um a um.",
            "Para responder bem, o assistente tem de perceber o vocabulário de cada pessoa: "
            "em que campo aparece «roupeiro», se «a obra está parada» é Desenho ou Produção, "
            "e se «o Silva» é um cliente ou um responsável.",
            "Ninguém conhece estas expressões melhor do que quem trabalha todos os dias com o "
            "Martelo. Escreve como falas, mesmo com abreviaturas ou nomes informais. O que "
            "interessa é que os exemplos sejam reais.",
        ],
    )

    caixa(
        doc,
        "O que o assistente faz — e o que não faz",
        [
            "FAZ: procura nas obras e nos orçamentos que já existem e mostra os resultados.",
            "FAZ: pergunta, quando uma expressão tiver mais do que um significado possível.",
            "NÃO FAZ: não cria, não altera nem apaga nada. Não mexe em estados, datas, preços "
            "ou obras.",
            "NÃO FAZ: não inventa valores. Tudo o que mostra é lido da base de dados.",
            "Se o resultado não for o esperado, ignoras e perguntas de outra maneira. O pior "
            "que pode acontecer é mostrar obras erradas — nunca estragar dados.",
        ],
    )

    caixa(
        doc,
        "Não é obrigatório — mas sem isto o assistente responde pior",
        [
            "Preencher este documento é voluntário e não há nada a controlar aqui.",
            "Só que o assistente aprende com o que estiver escrito: quem não der exemplos, vai "
            "ter respostas mais fracas às suas próprias perguntas. Vale a pena o quarto de "
            "hora.",
            "Não é preciso preencher tudo. O que souberes já ajuda, e podes deixar quadros em "
            "branco.",
        ],
        cor_borda="8A5A00",
        fundo=CREME_AVISO,
        cor_titulo=AVISO,
    )

    doc.add_page_break()

    # ---------------- parte 0 ----------------
    titulo(doc, "Parte 0 · Quem está a preencher")
    tabela(
        doc,
        ["Nome", "O que faz no dia a dia", "Data"],
        [4.5, 9.5, 3.0],
        fixos=[[nome_pessoa, "", ""]],
        vazias=0,
    )

    # ---------------- parte 1 ----------------
    titulo(doc, "Parte 1 · As perguntas que farias")
    paragrafo(
        doc,
        "Esta é a parte mais importante do documento. Escreve as perguntas tal como as dirias "
        "a um colega, sem te preocupares com a escrita nem com a forma como o computador vai "
        "perceber.",
    )
    paragrafo(
        doc,
        "Ao lado, escreve o resultado que esperavas ver. É isso que permite confirmar se o "
        "assistente acertou.",
    )
    paragrafo(
        doc,
        "As linhas a cinzento são exemplos reais, do documento já preenchido pelo Paulo. "
        "Acrescenta as tuas.",
        italic=True,
        cor=CINZENTO,
    )
    tabela(
        doc,
        ["Pergunta que eu faria", "O que eu esperava que aparecesse"],
        [8.5, 8.5],
        exemplos=[
            [
                "Que obras minhas estão atrasadas?",
                "As minhas obras com a Data Entrega já passada e que ainda não estão Arquivadas",
            ],
            [
                "Já fizemos algum closet para a J.F. Viva?",
                "Obras desse cliente com «closet» na descrição ou nas notas",
            ],
            [
                "Obras onde usei dobradiça XPTO",
                "Obras com essa referência ou marca nas notas ou nas matérias",
            ],
            [
                "Obras que estão há muito tempo em Desenho",
                "Obras ainda em Desenho, a assinalar que podem estar esquecidas",
            ],
        ],
        vazias=14,
    )

    doc.add_page_break()

    # ---------------- parte 2 ----------------
    titulo(doc, "Parte 2 · O vocabulário da casa")
    paragrafo(
        doc,
        "Aqui explicas o significado das palavras. Em cada quadro: à esquerda a expressão que "
        "costumas usar, à direita onde é que ela se encontra na obra. Se não souberes o nome "
        "do campo, escreve por palavras tuas.",
    )
    paragrafo(
        doc,
        "No fim do documento está a lista de todos os campos de uma obra, para consultares.",
        italic=True,
    )

    titulo(doc, "2.1 · Tipos de trabalho e de móvel", nivel=2)
    tabela(
        doc,
        ["Palavra que usamos", "Outras formas de dizer o mesmo", "Onde aparece na obra"],
        [5.0, 6.0, 6.0],
        exemplos=[
            ["roupeiro", "roupeiros, guarda-fatos", "Descrição produção"],
            ["canto cego", "canto feijão", "Descrição artigos, Notas"],
        ],
        vazias=10,
    )

    titulo(doc, "2.2 · Materiais e acabamentos", nivel=2)
    tabela(
        doc,
        ["Palavra que usamos", "O que significa exatamente", "Onde aparece na obra"],
        [5.0, 6.0, 6.0],
        exemplos=[
            ["lacado", "obra que leva lacagem, seja qual for a cor", "Descrição produção, Notas"],
            ["sandwich", "painel produzido, colagem, HPL, termolaminado", "Matérias usados"],
        ],
        vazias=8,
    )

    doc.add_page_break()

    titulo(doc, "2.3 · Estados da obra", nivel=2)
    paragrafo(
        doc,
        "No Martelo há quatro estados: Desenho, Produção, Finalizado e Arquivado. Na conversa "
        "dizem-se outras coisas. Escreve as expressões que usas e a que estado correspondem.",
    )
    caixa(
        doc,
        "Uma pergunta a que precisamos mesmo de resposta",
        [
            "Quando dizes que uma obra «já está fechada» ou «já está despachada» — isso é "
            "Finalizado ou Arquivado? Escreve no quadro qual dos dois usas.",
            "Isto é importante porque o Martelo deixa de dar alerta de atraso às obras "
            "Finalizadas e Arquivadas.",
        ],
        cor_borda="8A5A00",
        fundo=CREME_AVISO,
        cor_titulo=AVISO,
    )
    tabela(
        doc,
        ["Expressão que dizemos", "A que estado corresponde"],
        [8.5, 8.5],
        exemplos=[
            ["está na máquina", "Produção"],
            ["está parada, à espera de material", "Produção (mas quero saber que está parada)"],
        ],
        vazias=9,
    )

    titulo(doc, "2.4 · Pessoas", nivel=2)
    paragrafo(
        doc,
        "Os nomes da coluna da direita são os que estão escritos no Martelo — repara que são "
        "sem acentos. Na coluna da esquerda escreve como é que tratas cada pessoa na "
        "conversa: alcunhas, apelidos, nomes com acento, iniciais.",
    )
    paragrafo(
        doc,
        "Exemplo do que nos interessa: se disseres «a Elsa» e no Martelo estiver «Elisabete», "
        "o assistente tem de saber que é a mesma pessoa.",
        italic=True,
        cor=CINZENTO,
    )
    tabela(
        doc,
        ["Como lhe chamamos (escreve aqui)", "Nome que está no Martelo"],
        [8.5, 8.5],
        fixos=[["", nome] for nome in RESPONSAVEIS_REAIS],
        vazias=3,
    )

    doc.add_page_break()

    titulo(doc, "2.5 · Clientes", nivel=2)
    paragrafo(
        doc,
        "Abreviaturas e formas curtas dos clientes com que mais trabalhas. Muitas vezes é o "
        "nome simplex (o nome curto das pastas) que se usa a falar.",
    )
    tabela(
        doc,
        ["Como dizemos", "Nome completo do cliente"],
        [8.5, 8.5],
        exemplos=[["a Viva, o JF, JF_VIVA", "MÓVEIS J.F. VIVA"]],
        vazias=10,
    )

    titulo(doc, "2.6 · Tempo e urgência", nivel=2)
    paragrafo(
        doc,
        "Expressões de tempo, e o que significam para ti em dias concretos. É importante ser "
        "preciso: para uns «urgente» são 2 dias, para outros é esta semana.",
    )
    tabela(
        doc,
        ["Expressão", "O que significa em dias / datas", "Conta a partir de que data?"],
        [4.5, 7.0, 5.5],
        exemplos=[
            ["urgente", "entrega nos próximos 2 dias", "Data Entrega"],
            ["obra antiga", "começou há mais de 2 meses", "Data Início"],
        ],
        vazias=8,
    )

    titulo(doc, "2.7 · Outras palavras", nivel=2)
    tabela(
        doc,
        ["Expressão", "O que quero dizer com isso"],
        [6.0, 11.0],
        exemplos=[
            ["obra grande", "acima de 10 000 € ou com mais de 10 artigos"],
            ["assistência", "pedidos adicionais, retificações, correções de artigos"],
        ],
        vazias=8,
    )

    doc.add_page_break()

    # ---------------- parte 3 ----------------
    titulo(doc, "Parte 3 · Palavras que podem confundir")
    paragrafo(
        doc,
        "Palavras que significam coisas diferentes conforme quem fala. São estas que mais "
        "fazem o assistente errar — e são as que ele deve perguntar em vez de adivinhar.",
    )
    tabela(
        doc,
        ["Palavra", "Pode querer dizer isto…", "…ou isto", "O que ele deve perguntar"],
        [3.0, 4.7, 4.7, 4.6],
        exemplos=[
            [
                "Silva",
                "o cliente Silva",
                "o colega Pedro Silva",
                "«Silva é o cliente ou o responsável?»",
            ],
            [
                "portas",
                "portas de roupeiro",
                "portas de abrir ou de correr",
                "«São portas de roupeiro? De abrir ou de correr?»",
            ],
        ],
        vazias=7,
    )

    # ---------------- parte 4 ----------------
    titulo(doc, "Parte 4 · Como preferes a resposta")
    paragrafo(doc, "Assinala com um X a opção que preferes. Não há resposta certa.")
    tabela(
        doc,
        ["X", "Forma de responder"],
        [1.5, 15.5],
        exemplos=[
            ["", "Só a lista de obras, sem texto nenhum"],
            [
                "",
                "Uma frase curta a dizer o que percebeu («Encontrei 4 obras do Paulo em "
                "atraso») e a lista por baixo",
            ],
            ["", "Só texto, com os processos sublinhados para eu clicar"],
        ],
        vazias=2,
    )
    paragrafo(doc, "Outra ideia tua sobre a forma de responder:")
    tabela(doc, ["Escreve aqui"], [17.0], vazias=3)

    doc.add_page_break()

    # ---------------- parte 5 ----------------
    titulo(doc, "Parte 5 · Avisos que te dariam jeito")
    paragrafo(
        doc,
        "Além de responder a perguntas, o assistente pode acompanhar os dados e avisar-te. A "
        "ideia é avisar só quando for mesmo útil.",
    )
    tabela(
        doc,
        ["Aviso que me dava jeito", "De quanto em quanto tempo"],
        [11.5, 5.5],
        exemplos=[
            [
                "Orçamentos enviados há mais de 15 dias sem resposta do cliente",
                "Uma vez por semana, à segunda-feira",
            ],
            [
                "Obras finalizadas há muito tempo que o cliente ainda não levantou",
                "Uma vez por semana",
            ],
        ],
        vazias=8,
    )

    titulo(doc, "E o que NÃO queres ver", nivel=2)
    caixa(
        doc,
        "Este quadro conta tanto como o de cima — por favor não o deixes em branco",
        [
            "Escreve aqui tudo o que te iria irritar, interromper ou dar a sensação de estares "
            "a ser vigiado.",
            "Se disseres que não queres contagens do teu trabalho, ou avisos a meio do dia, "
            "isso passa a ser uma regra do programa. É a maneira de garantir que o assistente "
            "não se torna um chato.",
        ],
        cor_borda="8A5A00",
        fundo=CREME_AVISO,
        cor_titulo=AVISO,
    )
    tabela(
        doc,
        ["Isto não quero que apareça", "Porquê"],
        [8.5, 8.5],
        exemplos=[["Contagens de quantas obras fiz por semana", "Não quero sentir-me avaliado"]],
        vazias=6,
    )

    doc.add_page_break()

    # ---------------- anexo ----------------
    titulo(doc, "Anexo · Os campos que existem numa obra")
    paragrafo(
        doc,
        "Lista para consulta, para quando preencheres a coluna «onde aparece na obra». Não é "
        "preciso preencher nada aqui.",
    )

    caixa(
        doc,
        "Os seis campos onde está a descrição do trabalho",
        [
            "Descrição artigos · Matérias usados · Descrição produção · Notas 1, 2 e 3",
            "É nestes seis que está o resumo da obra: materiais, referências, ferragens, "
            "descrições, eletrodomésticos e artigos especiais. É aqui que a pesquisa vai "
            "procurar quase sempre.",
        ],
    )

    for nome_grupo, linhas in (
        (
            "Identificação",
            [
                ["Processo", "Código completo, por exemplo 26.1134_01_01_JF_VIVA"],
                ["Ano", "Ano da obra, por exemplo 2026"],
                ["N.º Enc. PHC", "Número da encomenda no PHC, por exemplo 1134"],
                ["V. Obra / V. CutRite", "Versões da obra e do plano de corte"],
                ["Criada em", "Data em que a obra entrou nesta lista"],
            ],
        ),
        (
            "Cliente e orçamento",
            [
                ["Cliente", "Nome completo, por exemplo MÓVEIS J.F. VIVA"],
                ["Cliente simplex", "Nome curto usado nas pastas, por exemplo JF_VIVA"],
                ["N.º Cliente PHC", "Número do cliente no PHC"],
                ["Ref. Cliente", "Referência fornecida pelo cliente"],
                ["N.º Orçamento / V. Orç.", "Orçamento de origem e a sua versão"],
                ["Preço total", "Valor da obra"],
                ["Qt. artigos", "Quantidade de artigos da obra"],
            ],
        ),
        (
            "Estado e datas",
            [
                ["Estado", "Desenho, Produção, Finalizado ou Arquivado"],
                ["Responsável", "Pessoa que trata da obra"],
                ["Data Início", "Quando a obra começou"],
                ["Data Entrega", "Quando deve estar pronta"],
                ["Localização", "Onde a obra vai ser montada"],
                ["Obra", "Nome curto da obra"],
                ["Tipo Pasta", "Encomenda de Cliente ou Encomenda de Cliente Final"],
            ],
        ),
        (
            "Texto livre",
            [
                ["Descrição artigos", "Vem da proposta ou da encomenda importada do PHC"],
                ["Matérias usados", "Materiais, por exemplo AGL_MLM_LINHO_CANCUN_10/16/19MM"],
                ["Descrição produção", "O que se vai produzir, por exemplo 3 ROUPEIROS PORTAS ABRIR"],
                ["Notas 1, 2 e 3", "Observações escritas à mão, ferragens, avisos"],
            ],
        ),
    ):
        titulo(doc, nome_grupo, nivel=2)
        tabela(doc, ["Campo", "O que contém"], [5.5, 11.5], fixos=linhas, vazias=0)

    caixa(
        doc,
        "Quando acabares",
        [
            "Devolve o documento ao Paulo, por email ou em papel.",
            "Se te lembrares de mais alguma coisa depois de entregares, diz na mesma. Este "
            "perfil vai sendo melhorado à medida que formos usando o assistente.",
        ],
    )

    doc.save(str(destino))


if __name__ == "__main__":
    pasta = Path(sys.argv[1])
    pasta.mkdir(parents=True, exist_ok=True)
    for nome in sys.argv[2:]:
        destino = pasta / f"Ensinar_assistente_Martelo_{nome}.docx"
        gerar(destino, nome)
        print("escrito:", destino.name)
